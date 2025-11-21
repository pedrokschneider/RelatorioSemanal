"""
Gerador de relatórios otimizado combinando recursos do SimpleReportGenerator original e V2.
"""

import os
import sys
import inspect
import logging
import pandas as pd
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any

from ..config import ConfigManager

logger = logging.getLogger("ReportSystem")

# Função utilitária para parse de datas flexível
from datetime import datetime

def parse_data_flex(date_str):
    if not date_str:
        return None
    
    # Se já for um objeto datetime, retornar diretamente
    if hasattr(date_str, 'strftime'):
        return date_str
    
    # Se for string vazia ou None, retornar None
    if not isinstance(date_str, str) or not date_str.strip():
        return None
    
    # Lista de formatos para tentar
    formatos = [
        "%d/%m/%Y",      # 25/12/2024
        "%d/%m/%y",      # 25/12/24
        "%Y-%m-%d",      # 2024-12-25
        "%d/%m",         # 25/12
        "%Y-%m-%d %H:%M:%S",  # 2024-12-25 10:30:00
        "%d/%m/%Y %H:%M:%S",  # 25/12/2024 10:30:00
        "%Y-%m-%dT%H:%M:%S",  # 2024-12-25T10:30:00
        "%Y-%m-%dT%H:%M:%S.%f",  # 2024-12-25T10:30:00.123456
        "%Y-%m-%dT%H:%M:%S.%fZ", # 2024-12-25T10:30:00.123456Z
        "%Y-%m-%dT%H:%M:%SZ",   # 2024-12-25T10:30:00Z
    ]
    
    for fmt in formatos:
        try:
            return datetime.strptime(date_str, fmt)
        except Exception:
            continue
    
    # Se nenhum formato funcionou, tentar extrair apenas a parte da data
    import re
    # Padrão para extrair data no formato YYYY-MM-DD
    match = re.search(r'(\d{4})-(\d{2})-(\d{2})', date_str)
    if match:
        try:
            return datetime(int(match.group(1)), int(match.group(2)), int(match.group(3)))
        except Exception:
            pass
    
    # Padrão para extrair data no formato DD/MM/YYYY
    match = re.search(r'(\d{1,2})/(\d{1,2})/(\d{4})', date_str)
    if match:
        try:
            return datetime(int(match.group(3)), int(match.group(2)), int(match.group(1)))
        except Exception:
            pass
    
    # Padrão para extrair data no formato DD/MM
    match = re.search(r'(\d{1,2})/(\d{1,2})', date_str)
    if match:
        try:
            # Assumir ano atual
            current_year = datetime.now().year
            return datetime(current_year, int(match.group(2)), int(match.group(1)))
        except Exception:
            pass
    
    logger.warning(f"Não foi possível converter a data: '{date_str}'")
    return None

# 1. Ajustar formatação das listas (sem asteriscos, data sempre dd/mm)
def format_task_line(date_value, discipline, name, responsible=None):
    # Tenta converter para datetime
    dt = None
    if isinstance(date_value, str):
        dt = parse_data_flex(date_value)
        if not dt and len(date_value) >= 10:
            try:
                dt = datetime.strptime(date_value[:10], "%Y-%m-%d")
            except Exception:
                pass
    elif hasattr(date_value, 'strftime'):
        dt = date_value
    if dt:
        formatted_date = dt.strftime("%d/%m")
    else:
        formatted_date = str(date_value)[:5]  # fallback
    line = f"{formatted_date} | {discipline}:\n{name}"
    if responsible:
        line += f" (Responsável: {responsible})"
    return line

class SimpleReportGenerator:
    """Gera relatórios com formato personalizado e links para o Construflow."""
    
    def __init__(self, config: ConfigManager):
        """
        Inicializa o gerador de relatórios.
        
        Args:
            config: Instância do ConfigManager
        """
        self.config = config
        self.prompt_template = self._load_prompt_template()
        # Criar diretório de relatórios se não existir
        self.reports_dir = os.path.join(os.getcwd(), "reports")
        os.makedirs(self.reports_dir, exist_ok=True)
    
    def _load_prompt_template(self) -> str:
        """Carrega o template de prompt embutido no código, ignorando arquivo externo."""
        return """Bom dia a todos,

Segue o resumo semanal do projeto [NOME_PROJETO].

---

🛎️ Pontos que precisam de resposta ([NOME_CLIENTE]):
[APONTAMENTOS_CLIENTE]

---

✅ Realizados na semana:
[TAREFAS_REALIZADAS]

---

📅 Planejamento para a próxima semana (atividades a iniciar):
[ATIVIDADES_INICIADAS_PROXIMA_SEMANA]

---

⚠️ Atrasos e desvios do período:
[ATRASOS_PERIODO]

---

📦 Entregas previstas para as próximas semanas:
[PROGRAMACAO_SEMANA]

---

 📊 Apontamentos pendentes por disciplina:
[TABELA_APONTAMENTOS]

Qualquer dúvida, estamos à disposição!
"""
    
    def generate_report(self, data: Dict[str, Any]) -> str:
        """
        Gera um relatório para o projeto usando template personalizado.
        
        Args:
            data: Dados processados do projeto
            
        Returns:
            Texto do relatório gerado
        """
        if not self.prompt_template:
            logger.error("Template de prompt não disponível")
            return "Erro: Template de relatório não disponível."
        
        # Validar que data não é None
        if data is None:
            logger.error("Dados do projeto são None - não é possível gerar relatório")
            return "Erro: Dados do projeto não disponíveis."
        
        # Obter nome do cliente corretamente
        project_name = data.get('project_name', 'Projeto') if isinstance(data, dict) else 'Projeto'
        project_id = data.get('project_id', '') if isinstance(data, dict) else ''
        system = self._get_system_instance()  # Obtém instância do sistema
        client_names = system.get_client_names(project_id) if system else []
        if client_names and client_names[0]:
            client_name = client_names[0]
        else:
            client_name = "Cliente"
        
        # Preparar substituições básicas
        replacements = {
            "[NOME_PROJETO]": project_name,
            "[NOME_CLIENTE]": client_name,
            "[DATA_ATUAL]": datetime.now().strftime("%d/%m/%Y")

        }
        
        # Gerar apontamentos do cliente com links
        apontamentos_cliente = self._gerar_apontamentos_cliente(data)
        replacements["[APONTAMENTOS_CLIENTE]"] = apontamentos_cliente
        
        # Gerar lista de tarefas realizadas - garantir funcionamento com os dois placeholders
        tarefas_realizadas = self._gerar_tarefas_realizadas(data)
        replacements["[TAREFAS_REALIZADAS]"] = tarefas_realizadas
        replacements["[TAREFAS_CONCLUIDAS]"] = tarefas_realizadas
        
        # Gerar lista de atrasos do período - garantir funcionamento com os dois placeholders
        atrasos_periodo = self._gerar_atrasos_periodo(data)
        replacements["[ATRASOS_PERIODO]"] = atrasos_periodo
        replacements["[TAREFAS_ATRASADAS]"] = atrasos_periodo
        
        # Gerar atividades que irão iniciar na próxima semana
        atividades_iniciadas_proxima_semana = self._gerar_atividades_iniciadas_proxima_semana(data)
        replacements["[ATIVIDADES_INICIADAS_PROXIMA_SEMANA]"] = atividades_iniciadas_proxima_semana
        
        # Gerar programação da semana - garantir funcionamento com os dois placeholders
        programacao_semana = self._gerar_programacao_semana(data)
        replacements["[PROGRAMACAO_SEMANA]"] = programacao_semana
        replacements["[TAREFAS_PROGRAMADAS]"] = programacao_semana
        
        # Gerar tabela de apontamentos por disciplina
        tabela_apontamentos = self._gerar_tabela_apontamentos(data)
        replacements["[TABELA_APONTAMENTOS]"] = tabela_apontamentos
        
        # Aplicar substituições
        report = self.prompt_template
        for key, value in replacements.items():
            report = report.replace(key, str(value))
        
        return report
    
    def _get_system_instance(self):
        """
        Tenta obter uma instância do sistema WeeklyReportSystem.
        Necessário para acessar métodos como get_client_names.
        """
        # Procurar classes WeeklyReportSystem no módulo atual
        for name, obj in inspect.getmembers(sys.modules['__main__']):
            if inspect.isclass(obj) and name == 'WeeklyReportSystem':
                # Procurar instâncias dessa classe
                for frame in inspect.stack():
                    for var in frame[0].f_locals.values():
                        if isinstance(var, obj):
                            return var
        
        # Se não encontrou, verificar se self.config é da classe ConfigManager
        # e se foi passado por um WeeklyReportSystem
        return None
    
    def _gerar_apontamentos_cliente(self, data: Dict[str, Any]) -> str:
        """Gera a seção de apontamentos que precisam de resposta do cliente."""
        # Validar que data não é None
        if data is None or not isinstance(data, dict):
            logger.warning("Dados são None ou inválidos em _gerar_apontamentos_cliente")
            return "Sem apontamentos pendentes para o cliente nesta semana."
        
        project_id = data.get('project_id', '')
        
        # Verificar se temos dados de apontamentos
        construflow_data = data.get('construflow_data')
        if not construflow_data or not construflow_data.get('active_issues'):
            return "Sem apontamentos pendentes para o cliente nesta semana."
            
        # Obter apontamentos do cliente
        system = self._get_system_instance()
        active_issues = data.get('construflow_data', {}).get('active_issues', [])
        client_issues = data.get('construflow_data', {}).get('client_issues', [])
        
        # Usar client_issues se disponível, caso contrário tentar filtrar active_issues
        if client_issues:
            logger.info(f"Usando {len(client_issues)} apontamentos já filtrados do cliente")
        else:
            # Tentar usar o método filter_client_issues do sistema
            try:
                import pandas as pd
                if system and hasattr(system, 'filter_client_issues') and active_issues:
                    df_issues = pd.DataFrame(active_issues)
                    filtered_issues = system.filter_client_issues(df_issues, project_id).to_dict('records')
                    client_issues = filtered_issues
                    logger.info(f"Filtrados {len(client_issues)} apontamentos via system.filter_client_issues")
            except Exception as e:
                logger.warning(f"Erro ao filtrar apontamentos do cliente: {e}")
                # Fallback: usar os primeiros 5 active_issues
                client_issues = active_issues[:5]
                logger.info(f"Usando {len(client_issues)} primeiros apontamentos como fallback")
        
        if not client_issues:
            return "Sem apontamentos pendentes para o cliente nesta semana."
        
        # Filtrar apontamentos - CORRIGIDO: apenas issues com status 'todo' na disciplina
        # CORREÇÃO: Incluir apenas issues com status da disciplina = 'todo'
        todo_issues = []
        

        
        for issue in client_issues:
            # CORREÇÃO: Incluir apenas issues com status da disciplina = 'todo'
            # Não incluir issues com status 'done' ou 'follow'
            if issue.get('status_y') == 'todo':
                todo_issues.append(issue)
        
        logger.info(f"Filtrados {len(todo_issues)} apontamentos com status 'todo' de {len(client_issues)} apontamentos do cliente")
        
        # Se não houver apontamentos com status 'todo', retornar mensagem informativa
        if not todo_issues:
            logger.info("Nenhum apontamento com status 'todo' encontrado")
            return "Sem apontamentos pendentes (A Fazer) para o cliente nesta semana."
        
        # Obter todas as issues do projeto do cache (issues_cache)
        all_issues = data.get('construflow_data', {}).get('all_issues', [])
        issues_cache = {}
        
        # Criar um dicionário de issues por code para busca rápida
        # Isso é necessário para garantir que o link correto para o apontamento seja gerado
        # usando o ID do apontamento que está em issues_cache, não em issues-disciplines_cache
        for issue in all_issues:
            # O projectId pode não estar disponível após o merge, então confiamos que
            # todas as issues em all_issues são do projeto atual
            if issue.get('code'):
                issues_cache[str(issue.get('code'))] = issue
        
        logger.info(f"Encontradas {len(issues_cache)} issues no cache para busca por code")
        
        # Agora vamos buscar nos dados originais de issues para garantir que o ID correto seja obtido
        raw_issues = {}  # Inicializar a variável
        try:
            # Tentar obter diretamente do connector
            if hasattr(self, 'construflow') and self.construflow:
                # Usar o método correto do connector GraphQL
                if hasattr(self.construflow, 'get_project_issues'):
                    issues_df = self.construflow.get_project_issues(project_id)
                elif hasattr(self.construflow, 'get_issues'):
                    issues_df = self.construflow.get_issues()
                else:
                    issues_df = None
            elif system and hasattr(system, 'processor') and hasattr(system.processor, 'construflow'):
                construflow_connector = system.processor.construflow
                if hasattr(construflow_connector, 'get_project_issues'):
                    issues_df = construflow_connector.get_project_issues(project_id)
                elif hasattr(construflow_connector, 'get_issues'):
                    issues_df = construflow_connector.get_issues()
                else:
                    issues_df = None
            else:
                issues_df = None
            
            if issues_df is not None and not issues_df.empty:
                # Converter para dicionário para busca rápida
                # Chave é uma tupla (project_id, code)
                for _, row in issues_df.iterrows():
                    if pd.notna(row.get('code')) and pd.notna(row.get('projectId')):
                        key = (str(row['projectId']), str(row['code']))
                        raw_issues[key] = row.to_dict()
                
                logger.info(f"Carregadas {len(raw_issues)} issues brutas para busca precisa por (projectId, code)")
        except Exception as e:
            logger.warning(f"Erro ao carregar issues brutas: {e}")
            # raw_issues já foi inicializado como {}
        
        # Agrupar issues por prioridade E por disciplina
        # Estrutura: issues_por_prioridade[prioridade][disciplina] = [issues]
        issues_por_prioridade = {
            'alta': {},
            'media': {},
            'baixa': {},
            'sem_prioridade': {}
        }
        
        # Obter disciplinas do cliente para agrupamento
        disciplinas_cliente = []
        try:
            if system and hasattr(system, 'get_client_disciplines'):
                disciplinas_cliente = system.get_client_disciplines(project_id)
                logger.info(f"Disciplinas do cliente para agrupamento: {disciplinas_cliente}")
        except Exception as e:
            logger.warning(f"Erro ao obter disciplinas do cliente para agrupamento: {e}")
        
        # Determinar se devemos mostrar separação por disciplinas
        # Mostrar separação se houver mais de uma disciplina do cliente
        mostrar_separacao_disciplinas = len(disciplinas_cliente) > 1
        
        # Processar cada issue para encontrar o ID correto e criar o link
        for issue in todo_issues:
            issue_code = str(issue.get('code', 'N/A'))
            issue_title = issue.get('title', 'Apontamento sem título')
            issue_disciplina = issue.get('name', '')  # Campo 'name' contém o nome da disciplina
            

            
            # Buscar o ID correto primeiro no dicionário de issues brutas usando (project_id, code)
            correct_issue_id = None
            
            # 1. Primeiro tentar encontrar o ID exato nas issues brutas (melhor opção)
            if raw_issues and (str(project_id), issue_code) in raw_issues:
                correct_issue_id = raw_issues[(str(project_id), issue_code)]['id']
                priority_raw = raw_issues[(str(project_id), issue_code)].get('priority')
                logger.info(f"Encontrado ID direto por (projectId, code) para o apontamento {issue_code}: {correct_issue_id}")
            # 2. Se não encontrar, tentar no issues_cache (pode ser menos preciso devido ao merge)
            elif issue_code in issues_cache and issues_cache[issue_code].get('id'):
                correct_issue_id = issues_cache[issue_code].get('id')
                priority_raw = issues_cache[issue_code].get('priority')
                logger.info(f"Encontrado ID via issues_cache para o apontamento {issue_code}: {correct_issue_id}")
            # 3. Fallback para o ID atual (antigo)
            else:
                correct_issue_id = issue.get('id')
                priority_raw = issue.get('priority')
                if correct_issue_id:
                    logger.warning(f"Não foi encontrado ID para o apontamento {issue_code}. Usando ID atual: {correct_issue_id}")
                else:
                    # Se nem o ID antigo estiver disponível, usar um ID genérico (não ideal, mas evita links quebrados)
                    correct_issue_id = "0"
                    logger.error(f"Não foi possível determinar o ID para o apontamento {issue_code}. Usando ID genérico.")
            
            # Construir o link para o apontamento com o ID correto
            construflow_url = f"https://app.construflow.com.br/workspace/project/{project_id}/issues?issueId={correct_issue_id}"
        
            # Calcular tempo desde a última atualização
            dias_sem_atualizacao = ""
            updated_at = issue.get('updatedAt')
            if updated_at:
                try:
                    from datetime import datetime
                    if isinstance(updated_at, str):
                        if 'Z' in updated_at:
                            updated_at = updated_at.replace('Z', '+00:00')
                        updated_date = datetime.fromisoformat(updated_at)
                    else:
                        updated_date = updated_at
                    agora = datetime.now()
                    if updated_date.tzinfo:
                        agora = agora.replace(tzinfo=updated_date.tzinfo)
                    diferenca = agora - updated_date
                    dias = diferenca.days
                    if dias == 0:
                        dias_sem_atualizacao = "(sem atualização hoje)"
                    elif dias == 1:
                        dias_sem_atualizacao = "(sem atualização há 1 dia)"
                    else:
                        dias_sem_atualizacao = f"(sem atualização há {dias} dias)"
                except Exception:
                    pass
            
            # Determinar em qual grupo de prioridade colocar essa issue
            priority = str(priority_raw).lower() if priority_raw else None
            
            # Mapeamento de valores possíveis de prioridade
            if priority in ['high', 'alta', '3', 3]:
                priority_group = 'alta'
            elif priority in ['medium', 'media', 'média', '2', 2]:
                priority_group = 'media'
            elif priority in ['low', 'baixa', '1', 1]:
                priority_group = 'baixa'
            else:
                priority_group = 'sem_prioridade'
            
            # Determinar a disciplina da issue
            # Normalizar o nome da disciplina para comparação case-insensitive
            disciplina_normalized = str(issue_disciplina).strip().lower() if issue_disciplina else ''
            
            # Usar o nome real da disciplina do cliente, não "Cliente 01", "Cliente 02"
            disciplina_key = 'outras'
            if disciplinas_cliente and issue_disciplina:
                # Procurar a disciplina correspondente na lista do cliente
                for disc_cliente in disciplinas_cliente:
                    if disc_cliente and str(disc_cliente).strip().lower() == disciplina_normalized:
                        # Usar o nome real da disciplina como está configurado
                        disciplina_key = str(disc_cliente).strip()
                        break
                # Se não encontrou nas disciplinas do cliente, usar o nome da disciplina da issue
                if disciplina_key == 'outras' and issue_disciplina:
                    disciplina_key = str(issue_disciplina).strip()
            elif issue_disciplina:
                # Se não há disciplinas do cliente configuradas, usar o nome da disciplina da issue
                disciplina_key = str(issue_disciplina).strip()
            
            # Inicializar o dicionário da disciplina se não existir
            if disciplina_key not in issues_por_prioridade[priority_group]:
                issues_por_prioridade[priority_group][disciplina_key] = []
            
            # Armazenar a linha formatada no grupo correto (prioridade + disciplina)
            if dias_sem_atualizacao and 'sem atualização' in dias_sem_atualizacao:
                issue_line = f"[#{issue_code}]({construflow_url}) – {issue_title}\n   ⏳ {dias_sem_atualizacao.strip('()')}"
            else:
                issue_line = f"[#{issue_code}]({construflow_url}) – {issue_title}"
            issues_por_prioridade[priority_group][disciplina_key].append(issue_line)
        
        # Construir o resultado final agrupado por prioridade E por disciplina
        result = ""
        
        # Função auxiliar para renderizar uma seção de prioridade com separação por disciplinas
        # Usar mostrar_separacao_disciplinas do escopo externo
        def render_priority_section(priority_group, emoji, title, mostrar_sep=mostrar_separacao_disciplinas):
            """Renderiza uma seção de prioridade com separação por disciplinas."""
            section_result = ""
            if issues_por_prioridade[priority_group]:
                # Verificar se há issues em alguma disciplina
                has_issues = any(issues_por_prioridade[priority_group].values())
                if has_issues:
                    section_result += f"{emoji} {title}\n"
                    
                    # Ordenar disciplinas conforme a ordem configurada nas disciplinas do cliente
                    # Manter a ordem das disciplinas do cliente
                    disciplinas_ordenadas = list(issues_por_prioridade[priority_group].keys())
                    
                    # Se temos disciplinas do cliente configuradas, ordenar por essa ordem
                    if disciplinas_cliente:
                        # Criar um dicionário de ordem baseado nas disciplinas do cliente
                        ordem_disciplinas = {str(d).strip().lower(): idx for idx, d in enumerate(disciplinas_cliente)}
                        
                        # Ordenar: primeiro as que estão na lista do cliente (na ordem configurada), depois outras
                        disciplinas_ordenadas = sorted(
                            disciplinas_ordenadas,
                            key=lambda x: (
                                ordem_disciplinas.get(str(x).strip().lower(), 999),  # Ordem na lista do cliente
                                x  # Nome como fallback
                            )
                        )
                    
                    for disciplina_key in disciplinas_ordenadas:
                        issues_da_disciplina = issues_por_prioridade[priority_group][disciplina_key]
                        if issues_da_disciplina:
                            # Adicionar cabeçalho da disciplina se houver múltiplas disciplinas do cliente
                            # Mostra o nome real da disciplina (ex: [Produto Planeta], [Projetos Planeta])
                            if mostrar_sep:
                                section_result += f"\n[{disciplina_key}]:\n\n"
                            
                            # Adicionar as issues da disciplina
                            for issue_line in issues_da_disciplina:
                                section_result += f"- {issue_line}\n"
                    
                    section_result += "\n"
            return section_result
        
        # Prioridade Alta - Vermelho
        result += render_priority_section('alta', '🔴', 'Prioridade Alta')
        
        # Prioridade Média - Laranja
        result += render_priority_section('media', '🟠', 'Prioridade Média')
        
        # Prioridade Baixa - Verde
        result += render_priority_section('baixa', '🟢', 'Prioridade Baixa')
        
        # Sem prioridade definida
        result += render_priority_section('sem_prioridade', '⚪', 'Sem Prioridade Definida')
        
        return result.strip() if result else "Sem apontamentos pendentes para o cliente nesta semana."

    def _is_status_done(self, status_raw: str) -> bool:
        if not status_raw:
            return False
        s = str(status_raw).strip().lower()
        # Smartsheet: valores exatos
        return s == 'feito'

    def _is_status_not_done(self, status_raw: str) -> bool:
        if not status_raw:
            return False
        s = str(status_raw).strip().lower()
        return s == 'não feito'

    def _gerar_tarefas_realizadas(self, data: Dict[str, Any]) -> str:
        """Gera a seção de tarefas realizadas no período."""
        if data is None or not isinstance(data, dict):
            logger.warning("Dados são None ou inválidos em _gerar_tarefas_realizadas")
            return "Sem tarefas realizadas no período."
        
        smartsheet_data = data.get('smartsheet_data', {})

        # Base de dados: usar all_tasks e excluir explicitamente Não Feito
        if isinstance(smartsheet_data, dict) and 'all_tasks' in smartsheet_data:
            all_tasks = smartsheet_data.get('all_tasks', [])
        elif isinstance(smartsheet_data, list):
            all_tasks = smartsheet_data
        else:
            logger.warning(f"Formato não reconhecido para smartsheet_data: {type(smartsheet_data)}")
            return "Sem tarefas concluídas no período."

        completed_tasks = []
        for task in all_tasks:
            if not isinstance(task, dict):
                continue
            status = task.get('Status')
            if self._is_status_not_done(status):
                continue
            if self._is_status_done(status):
                completed_tasks.append(task)
                continue
            # Se status não é fornecido claramente, não inferir "feito" por datas.

        if not completed_tasks:
            return "Sem tarefas concluídas no período."

        from datetime import datetime
        def get_task_date(task):
            task_date = task.get('Data Término', task.get('Data de Término', ''))
            if isinstance(task_date, str):
                for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d/%m"):
                    try:
                        return datetime.strptime(task_date, fmt)
                    except Exception:
                        continue
                return datetime.min
            elif hasattr(task_date, 'strftime'):
                return task_date
            return datetime.min

        completed_tasks.sort(key=get_task_date, reverse=True)

        tarefas_por_disciplina = {}
        for task in completed_tasks:
            task_date = task.get('Data Término', task.get('Data de Término', ''))
            task_name = task.get('Nome da Tarefa', task.get('Task Name', ''))
            task_discipline = task.get('Disciplina', task.get('Discipline', '')) or 'Sem Disciplina'

            dt = parse_data_flex(task_date)
            if not dt and isinstance(task_date, str) and len(task_date) >= 10:
                try:
                    so_data = task_date[:10]
                    dt = datetime.strptime(so_data, "%Y-%m-%d")
                except Exception:
                    pass
            if dt:
                formatted_date = dt.strftime("%d/%m")
            else:
                import re
                match = re.search(r'(\d{4})[-/](\d{2})[-/](\d{2})', str(task_date))
                if match:
                    formatted_date = f"{match.group(3)}/{match.group(2)}"
                else:
                    match2 = re.search(r'(\d{2})/(\d{2})', str(task_date))
                    if match2:
                        formatted_date = f"{match2.group(1)}/{match2.group(2)}"
                    else:
                        formatted_date = str(task_date).strip()[:5]

            linha = f"{formatted_date} │ {task_name}"
            tarefas_por_disciplina.setdefault(task_discipline, []).append(linha)

        if not tarefas_por_disciplina:
            return "Sem tarefas concluídas no período."

        result = ""
        for disciplina, tarefas in tarefas_por_disciplina.items():
            result += f"{disciplina}\n"
            for tarefa in tarefas:
                result += f"{tarefa}\n"
            result += "\n"
        return result.strip()

    def _gerar_atividades_iniciadas_proxima_semana(self, data: Dict[str, Any]) -> str:
        """
        Gera a seção de atividades que irão iniciar na próxima semana (segunda a domingo).
        Se o relatório for solicitado antes de sexta-feira, considera atividades da semana atual e da próxima.
        """
        if data is None or not isinstance(data, dict):
            logger.warning("Dados são None ou inválidos em _gerar_atividades_iniciadas_proxima_semana")
            return "Sem atividades programadas para iniciar na próxima semana."
        
        smartsheet_data = data.get('smartsheet_data', {})
        
        # Usar dados categorizados se disponíveis
        if isinstance(smartsheet_data, dict) and 'all_tasks' in smartsheet_data:
            all_tasks = smartsheet_data.get('all_tasks', [])
        elif isinstance(smartsheet_data, list):
            all_tasks = smartsheet_data
        else:
            logger.warning(f"Formato não reconhecido para smartsheet_data: {type(smartsheet_data)}")
            return "Sem atividades previstas para iniciar na próxima semana."
        
        logger.info(f"Processando {len(all_tasks)} tarefas para atividades que iniciam na próxima semana")
        
        from datetime import datetime, timedelta
        hoje = datetime.now()
        weekday = hoje.weekday()  # 0=segunda, 4=sexta
        
        # Calcular intervalo de datas
        if weekday < 4:  # Antes de sexta-feira
            # Segunda-feira desta semana
            segunda_atual = (hoje - timedelta(days=hoje.weekday())).replace(hour=0, minute=0, second=0, microsecond=0)
            # Domingo da próxima semana
            domingo_proxima = segunda_atual + timedelta(days=13)
            intervalo_inicio = segunda_atual
            intervalo_fim = domingo_proxima
            logger.info(f"Relatório antes de sexta-feira. Intervalo: {intervalo_inicio.strftime('%d/%m/%Y')} a {intervalo_fim.strftime('%d/%m/%Y')}")
        else:
            # Próxima segunda-feira
            dias_ate_segunda = (7 - hoje.weekday()) % 7 or 7
            proxima_segunda = (hoje + timedelta(days=dias_ate_segunda)).replace(hour=0, minute=0, second=0, microsecond=0)
            proximo_domingo = proxima_segunda + timedelta(days=6)
            intervalo_inicio = proxima_segunda
            intervalo_fim = proximo_domingo
            logger.info(f"Relatório após sexta-feira. Intervalo: {intervalo_inicio.strftime('%d/%m/%Y')} a {intervalo_fim.strftime('%d/%m/%Y')}")
        
        # Agrupar atividades por disciplina
        atividades_por_disciplina = {}
        tarefas_processadas = 0
        tarefas_com_data_inicio = 0
        tarefas_no_intervalo = 0
        
        for task in all_tasks:
            if not isinstance(task, dict):
                continue
            
            tarefas_processadas += 1
            data_inicio = task.get('Data Inicio', task.get('Data de Início', ''))
            data_termino = task.get('Data Término')
            
            # Log detalhado para debug
            task_name = task.get('Nome da Tarefa', task.get('Task Name', ''))
            
            if not data_inicio:
                logger.debug(f"Tarefa '{task_name}' sem data de início")
                continue
            
            tarefas_com_data_inicio += 1
            
            try:
                if isinstance(data_inicio, str):
                    data_inicio_dt = parse_data_flex(data_inicio)
                else:
                    data_inicio_dt = data_inicio
            except Exception as e:
                logger.warning(f"Erro ao processar data de início '{data_inicio}' para tarefa '{task_name}': {e}")
                continue
            
            if not data_inicio_dt:
                logger.debug(f"Tarefa '{task_name}' com data de início inválida: '{data_inicio}'")
                continue
            
            # Verificar se está no intervalo
            if intervalo_inicio <= data_inicio_dt <= intervalo_fim:
                tarefas_no_intervalo += 1
                logger.debug(f"Tarefa '{task_name}' com início em {data_inicio_dt.strftime('%d/%m/%Y')} está no intervalo")
                
                # Formatar datas SEM ANO
                data_inicio_fmt = data_inicio_dt.strftime("%d/%m")
                
                if data_termino:
                    try:
                        if isinstance(data_termino, str):
                            data_termino_dt = parse_data_flex(data_termino)
                        else:
                            data_termino_dt = data_termino
                        data_termino_fmt = data_termino_dt.strftime("%d/%m") if data_termino_dt else "?"
                    except Exception:
                        data_termino_fmt = str(data_termino)[:5]
                else:
                    data_termino_fmt = "?"
                
                nome = task.get('Nome da Tarefa', task.get('Task Name', ''))
                disciplina = task.get('Disciplina', task.get('Discipline', '')) or 'Sem Disciplina'
                
                # Linha agrupada
                if data_inicio_fmt == data_termino_fmt or not data_termino:
                    linha = f"{data_inicio_fmt} │ {nome}"
                else:
                    linha = f"{data_inicio_fmt} a {data_termino_fmt} │ {nome}"
                
                if disciplina not in atividades_por_disciplina:
                    atividades_por_disciplina[disciplina] = []
                atividades_por_disciplina[disciplina].append(linha)
            else:
                logger.debug(f"Tarefa '{task_name}' com início em {data_inicio_dt.strftime('%d/%m/%Y')} fora do intervalo")
        
        logger.info(f"Estatísticas: {tarefas_processadas} tarefas processadas, {tarefas_com_data_inicio} com data de início, {tarefas_no_intervalo} no intervalo")
        
        # Montar resultado agrupado
        if not atividades_por_disciplina:
            return "Sem atividades previstas para iniciar na próxima semana."
        
        result = ""
        for disciplina, atividades in atividades_por_disciplina.items():
            result += f"{disciplina}\n"
            for atividade in atividades:
                result += f"{atividade}\n"
            result += "\n"
        
        return result.strip()

    def _gerar_atrasos_periodo(self, data: Dict[str, Any]) -> str:
        """Gera a seção de atrasos e desvios do período, incluindo baseline e motivo de atraso."""
        if data is None or not isinstance(data, dict):
            logger.warning("Dados são None ou inválidos em _gerar_atrasos_periodo")
            return "Sem atrasos registrados no período."
        
        smartsheet_data = data.get('smartsheet_data', {})

        # Preferir lista já preparada pelo processador
        if isinstance(smartsheet_data, dict) and 'delayed_tasks' in smartsheet_data:
            delayed_tasks = smartsheet_data.get('delayed_tasks', [])
        else:
            # Fallback usando all_tasks
            all_tasks = smartsheet_data.get('all_tasks', []) if isinstance(smartsheet_data, dict) else (smartsheet_data if isinstance(smartsheet_data, list) else [])
            delayed_tasks = []
            for task in all_tasks:
                if not isinstance(task, dict):
                    continue
                status = task.get('Status')
                categoria_atraso = task.get('Categoria de atraso') or task.get('Delay Category')
                tem_categoria_atraso = categoria_atraso and str(categoria_atraso).strip() not in ['', 'nan', 'None']
                if self._is_status_not_done(status) or tem_categoria_atraso:
                    delayed_tasks.append(task)

        if not delayed_tasks:
            return "Não foram identificados atrasos no período."

        result = ""
        for task in delayed_tasks:
            if not isinstance(task, dict):
                continue
            task_discipline = task.get('Disciplina', task.get('Discipline', ''))
            task_name = task.get('Nome da Tarefa', task.get('Task Name', ''))
            nova_data = task.get('Data Término', task.get('Data de Término', task.get('Due Date', '')))

            import re
            baseline = task.get('Data de Fim - Baseline Otus')
            baseline_keys = [k for k in task.keys() if re.match(r'^Data de Fim - Baseline Otus R\\d+$', k)]
            if baseline_keys:
                baseline_keys.sort(key=lambda x: int(re.findall(r'R(\\d+)$', x)[0]), reverse=True)
                baseline = task.get(baseline_keys[0]) or baseline

            baseline_fmt = baseline if baseline else "-"
            motivo = task.get('Motivo de atraso')
            motivo_fmt = motivo if motivo else "-"

            status = str(task.get('Status', '')).strip().lower()
            if self._is_status_not_done(status) and (not motivo_fmt or motivo_fmt == "-"):
                motivo_fmt = "Tarefa não realizada (status: Não Feito)"

            nova_data_fmt = ""
            if nova_data:
                nova_data_dt = parse_data_flex(nova_data)
                if not nova_data_dt and hasattr(nova_data, 'strftime'):
                    nova_data_dt = nova_data
                if nova_data_dt:
                    nova_data_fmt = nova_data_dt.strftime("%d/%m")
                else:
                    if isinstance(nova_data, str) and len(nova_data) >= 10:
                        try:
                            so_data = nova_data[:10]
                            dt_tmp = datetime.strptime(so_data, "%Y-%m-%d")
                            nova_data_fmt = dt_tmp.strftime("%d/%m")
                        except Exception:
                            nova_data_fmt = str(nova_data)[:5]
                    else:
                        nova_data_fmt = str(nova_data)[:5]

            if baseline:
                baseline_dt = parse_data_flex(baseline)
                if baseline_dt:
                    baseline_fmt = baseline_dt.strftime("%d/%m")
                else:
                    baseline_fmt = str(baseline)[:5]
            else:
                baseline_fmt = "-"

            result += (f"* {task_discipline} – {task_name}\n"
                       f"    - Nova data programada: {nova_data_fmt}\n"
                       f"    - Data prevista inicial: {baseline_fmt}\n"
                       f"    - Motivo do atraso: {motivo_fmt}\n")
        return result if result else "Não foram identificados atrasos no período."
    
    def _gerar_programacao_semana(self, data: Dict[str, Any]) -> str:
        """Gera a seção de programação para as próximas duas semanas."""
        if data is None or not isinstance(data, dict):
            logger.warning("Dados são None ou inválidos em _gerar_programacao_semana")
            return "Sem programação disponível para as próximas semanas."
        
        smartsheet_data = data.get('smartsheet_data', {})
        
        # Usar dados categorizados se disponíveis
        if isinstance(smartsheet_data, dict) and 'scheduled_tasks' in smartsheet_data:
            future_tasks = smartsheet_data.get('scheduled_tasks', [])
            logger.info(f"Usando {len(future_tasks)} tarefas programadas categorizadas")
        elif isinstance(smartsheet_data, dict) and 'all_tasks' in smartsheet_data:
            all_tasks = smartsheet_data.get('all_tasks', [])
            from datetime import datetime, timedelta
            today = datetime.today()
            next_week_end = today + timedelta(days=14)
            future_tasks = []
            for task in all_tasks:
                if not isinstance(task, dict):
                    continue
                task_date = task.get('Data Término', task.get('Data de Término', task.get('Due Date', '')))
                if task_date:
                    if isinstance(task_date, str):
                        future_tasks.append(task)
                    else:
                        try:
                            if today < task_date <= next_week_end:
                                future_tasks.append(task)
                        except:
                            future_tasks.append(task)
            if not future_tasks and all_tasks:
                future_tasks = all_tasks[:min(3, len(all_tasks))]
            logger.info(f"Filtradas {len(future_tasks)} tarefas programadas de {len(all_tasks)} tarefas")
        elif isinstance(smartsheet_data, list):
            from datetime import datetime, timedelta
            today = datetime.today()
            next_week_end = today + timedelta(days=14)
            all_tasks = smartsheet_data
            future_tasks = []
            for task in all_tasks:
                if not isinstance(task, dict):
                    continue
                task_date = task.get('Data Término', task.get('Data de Término', task.get('Due Date', '')))
                if task_date:
                    if isinstance(task_date, str):
                        future_tasks.append(task)
                    else:
                        try:
                            if today < task_date <= next_week_end:
                                future_tasks.append(task)
                        except:
                            future_tasks.append(task)
            if not future_tasks and all_tasks:
                valid_tasks = [t for t in all_tasks if isinstance(t, dict)]
                future_tasks = valid_tasks[:min(3, len(valid_tasks))]
            logger.info(f"Filtradas {len(future_tasks)} tarefas programadas do formato antigo")
        else:
            logger.warning(f"Formato não reconhecido para smartsheet_data: {type(smartsheet_data)}")
            return "Sem atividades programadas para as próximas duas semanas."
        
        if not future_tasks:
            return "Sem atividades programadas para as próximas duas semanas."
        
        def get_task_date(task):
            task_date = task.get('Data Término', task.get('Data de Término', task.get('Due Date', '')))
            if isinstance(task_date, str):
                try:
                    return datetime.strptime(task_date, "%d/%m/%Y")
                except:
                    return datetime.now() + timedelta(days=14)
            return task_date if task_date else datetime.now() + timedelta(days=14)
        
        future_tasks.sort(key=get_task_date, reverse=False)
        
        # Agrupar entregas por disciplina
        entregas_por_disciplina = {}
        for task in future_tasks:
            if not isinstance(task, dict):
                continue
            task_date = task.get('Data Término', task.get('Data de Término', task.get('Due Date', '')))
            task_name = task.get('Nome da Tarefa', task.get('Task Name', ''))
            task_discipline = task.get('Disciplina', task.get('Discipline', '')) or 'Sem Disciplina'
            
            # Formatar data SEM negrito, sempre dd/mm
            dt = parse_data_flex(task_date)
            if not dt and isinstance(task_date, str) and len(task_date) >= 10:
                try:
                    so_data = task_date[:10]
                    dt = datetime.strptime(so_data, "%Y-%m-%d")
                except Exception:
                    pass
            if dt:
                formatted_date = dt.strftime("%d/%m")
            else:
                import re
                match = re.search(r'(\d{4})[-/](\d{2})[-/](\d{2})', str(task_date))
                if match:
                    formatted_date = f"{match.group(3)}/{match.group(2)}"
                else:
                    match2 = re.search(r'(\d{2})/(\d{2})', str(task_date))
                    if match2:
                        formatted_date = f"{match2.group(1)}/{match2.group(2)}"
                    else:
                        formatted_date = str(task_date).strip()[:5]
            
            linha = f"{formatted_date} │ {task_name}"
            if task_discipline not in entregas_por_disciplina:
                entregas_por_disciplina[task_discipline] = []
            entregas_por_disciplina[task_discipline].append(linha)
        
        # Montar resultado agrupado
        if not entregas_por_disciplina:
            return "Sem atividades programadas para as próximas duas semanas."
        
        result = ""
        for disciplina, entregas in entregas_por_disciplina.items():
            result += f"{disciplina}\n"
            for entrega in entregas:
                result += f"{entrega}\n"
            result += "\n"
        
        return result.strip()
        
    def _gerar_tabela_apontamentos(self, data: dict) -> str:
        """Gera uma tabela de apontamentos por disciplina mostrando apenas status 'todo' (A Fazer)."""
        # Validar que data não é None
        if data is None or not isinstance(data, dict):
            logger.warning("Dados são None ou inválidos em _gerar_tabela_apontamentos")
            return "Sem dados de apontamentos por disciplina."
        
        issues = data.get('construflow_data', {}).get('all_issues', []) if data.get('construflow_data') else []
        if not issues:
            return "Sem dados de apontamentos por disciplina."

        df = pd.DataFrame(issues)
        if df.empty or 'name' not in df.columns or 'status_y' not in df.columns:
            # Tentar alternativa da V2
            disciplines = data.get('construflow_data', {}).get('disciplines', {})
            if disciplines:
                # Formatar como tabela markdown
                discipline_rows = [
                    "| Disciplina | Quantidade de Apontamentos |",
                    "|------------|----------------------------|"
                ]
                
                for discipline, count in disciplines.items():
                    discipline_rows.append(f"| {discipline} | {count} |")
                
                return "\n".join(discipline_rows)
            
            return "Sem dados de apontamentos por disciplina."

        # Renomear colunas para facilitar
        df.rename(columns={
            'name': 'Disciplina',
            'status_y': 'Status'
        }, inplace=True)

        # Filtrar apenas os apontamentos com status 'todo'
        df_filtered = df[df['Status'] == 'todo']
        
        if df_filtered.empty:
            return "Sem apontamentos pendentes (A Fazer) no período."

        # Contar os apontamentos "A Fazer" por disciplina
        contagem_por_disciplina = df_filtered.groupby('Disciplina').size().reset_index(name='A Fazer')
        
        # Ordenar por quantidade de apontamentos (do maior para o menor)
        contagem_por_disciplina = contagem_por_disciplina.sort_values('A Fazer', ascending=False)
        
        # Criar a tabela markdown
        linhas = ["| Disciplina | A Fazer |",
                "|------------|---------|"]
        
        for _, row in contagem_por_disciplina.iterrows():
            linhas.append(f"| {row['Disciplina']} | {row['A Fazer']} |")
        
        return "\n".join(linhas)

    def save_report(self, report_text: str, project_name: str, 
                   format_type: str = 'md') -> str:
        """
        Salva o relatório em um arquivo local.
        
        Args:
            report_text: Texto do relatório
            project_name: Nome do projeto
            format_type: Formato do arquivo ('docx', 'txt' ou 'md')
            
        Returns:
            Caminho do arquivo salvo
        """        
        # Formatar nome do arquivo
        today_str = datetime.now().strftime("%Y-%m-%d")
        safe_project_name = ''.join(c if c.isalnum() or c in ' -_' else '_' for c in project_name)
        safe_project_name = safe_project_name.replace(" ", "_")
        
        # Priorizar formato MD para melhor compatibilidade com Google Docs
        format_type = str(format_type or 'md')
        if format_type.lower() == 'md' or format_type.lower() == 'markdown':
            # Salvar como arquivo markdown
            file_name = f"Relatorio_{safe_project_name}_{today_str}.md"
            file_path = os.path.join(self.reports_dir, file_name)
            
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(report_text)
                logger.info(f"Relatório salvo como MD em {file_path}")
                return file_path
            except Exception as e:
                logger.error(f"Erro ao salvar relatório MD: {e}")
                
                # Tentar salvar em um local alternativo
                alt_path = os.path.join(os.getcwd(), file_name)
                try:
                    with open(alt_path, 'w', encoding='utf-8') as f:
                        f.write(report_text)
                    logger.info(f"Relatório MD salvo em local alternativo: '{alt_path}'")
                    return alt_path
                except Exception as e2:
                    logger.error(f"Erro ao salvar relatório MD em local alternativo: {e2}")
                    # Continuar com outros formatos
        
        # Formato DOCX se solicitado
        if format_type.lower() == 'docx':
            try:
                from docx import Document
                from docx.shared import RGBColor, Pt
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                import re
                doc = Document()
                # Adicionar título principal
                title = doc.add_heading(f"Relatório Semanal - {project_name}", level=1)
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # Quebrar o relatório em linhas
                paragraphs = report_text.split('\n')
                current_para = None
                # Mapear seções para Heading 2
                secoes_h2 = [
                    "Pontos que precisam de resposta",
                    "Realizados na semana:",
                    "Planejamento para a próxima semana (atividades a iniciar):",
                    "Atrasos e desvios do período:",
                    "Entregas previstas para as próximas semanas:",
                    "Apontamentos pendentes por disciplina:",
                    "🔴 Prioridade Alta",
                    "🟠 Prioridade Média",
                    "🟢 Prioridade Baixa",
                    "⚪ Sem Prioridade Definida"
                ]
                cor_prioridade = {
                    "🔴 Prioridade Alta": RGBColor(255, 0, 0),
                    "🟠 Prioridade Média": RGBColor(255, 140, 0),
                    "🟢 Prioridade Baixa": RGBColor(0, 180, 0),
                    "⚪ Sem Prioridade Definida": RGBColor(120, 120, 120)
                }
                for line in paragraphs:
                    l_strip = line.strip()
                    # Título de saudação
                    if l_strip == "Bom dia a todos,":
                        doc.add_paragraph(l_strip)
                        continue
                    # Título de fechamento
                    if l_strip == "Qualquer dúvida, estamos à disposição!":
                        doc.add_paragraph(l_strip)
                        continue
                    # Seção principal (Heading 2) ou prioridade
                    if l_strip in secoes_h2:
                        heading = doc.add_heading(l_strip, level=2)
                        if l_strip in cor_prioridade:
                            for run in heading.runs:
                                run.font.color.rgb = cor_prioridade[l_strip]
                        current_para = None
                        continue
                    # Linhas normais, listas, etc
                    if l_strip.startswith('- '):
                        item_para = doc.add_paragraph()
                        item_para.style = 'List Bullet'
                        item_para.add_run(l_strip[2:])
                    elif l_strip.startswith('* '):
                        item_para = doc.add_paragraph()
                        item_para.style = 'List Bullet'
                        item_para.add_run(l_strip[2:])
                    elif l_strip == '---' or l_strip == '':
                        continue
                    else:
                        doc.add_paragraph(l_strip)
                file_name = f"Relatorio_{safe_project_name}_{today_str}.docx"
                file_path = os.path.join(self.reports_dir, file_name)
                doc.save(file_path)
                logger.info(f"Relatório salvo como DOCX em {file_path}")
                return file_path
            except ImportError:
                logger.warning("Módulo python-docx não encontrado. Salvando como TXT.")
                format_type = 'txt'
        
        # Padrão - salvar como TXT se nenhum dos outros formatos funcionou
        file_name = f"Relatorio_{safe_project_name}_{today_str}.txt"
        file_path = os.path.join(self.reports_dir, file_name)
        
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(report_text)
            logger.info(f"Relatório salvo como TXT em {file_path}")
            return file_path
        except Exception as e:
            logger.error(f"Erro ao salvar relatório TXT: {e}")
            
            # Tentar salvar em um local alternativo
            alt_path = os.path.join(os.getcwd(), file_name)
            try:
                with open(alt_path, 'w', encoding='utf-8') as f:
                    f.write(report_text)
                logger.info(f"Relatório TXT salvo em local alternativo: '{alt_path}'")
                return alt_path
            except Exception as e2:
                logger.error(f"Erro ao salvar relatório TXT em local alternativo: {e2}")
                return ""